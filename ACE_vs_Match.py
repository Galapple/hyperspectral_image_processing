import os
import spectral.io.envi as envi
from tabulate import tabulate
from scipy import integrate
from scipy.ndimage import convolve
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import pickle
from datetime import datetime


def fix_corrupted_spectrum(corrupted_spectrum):
    """
    Fixes a corrupted spectrum by replacing zero values with the average of their four neighbors.

    Parameters:
    - corrupted_spectrum: A 1D numpy array where 10% of its elements are set to zero.

    Returns:
    - A numpy array where zero values have been replaced by the average of their four neighbors.
    """
    # Length of the spectrum
    num_elements = len(corrupted_spectrum)

    # Create a copy of the spectrum to modify
    fixed_spectrum = np.copy(corrupted_spectrum)

    # Iterate over each element of the spectrum
    for i in range(num_elements):
        # Check if the current element is zero
        if fixed_spectrum[i] == 0:
            # Initialize neighbors list
            neighbors = []

            # Check each of the four potential neighbors (two on each side)
            if i > 0:  # Left neighbor
                neighbors.append(corrupted_spectrum[i - 1])
            if i < num_elements - 1:  # Right neighbor
                neighbors.append(corrupted_spectrum[i + 1])
            if i > 1:  # Second left neighbor
                neighbors.append(corrupted_spectrum[i - 2])
            if i < num_elements - 2:  # Second right neighbor
                neighbors.append(corrupted_spectrum[i + 2])

            # Calculate the average of the neighbors if any neighbors exist
            if neighbors:
                fixed_spectrum[i] = np.mean([n for n in neighbors if n != 0])

    return fixed_spectrum


def corrupt_spectrum(pixel_spectrum):
    """
    Randomly sets 10% of the elements in a pixel spectrum to zero.

    Parameters:
    - pixel_spectrum: A 1D numpy array representing the pixel spectrum.

    Returns:
    - A numpy array with 10% of its elements set to zero at random positions.
    """
    # Calculate the number of elements to set to zero
    num_elements = len(pixel_spectrum)
    num_to_zero = int(num_elements * 0.1)  # 10% of the elements

    # Generate random indices to set to zero
    indices_to_zero = np.random.choice(num_elements, num_to_zero, replace=False)

    # Create a copy of the pixel spectrum to modify
    corrupted_spectrum = np.copy(pixel_spectrum)

    # Set the chosen indices to zero
    corrupted_spectrum[indices_to_zero] = 0

    return corrupted_spectrum


def compute_cov_matrix(hyper_cube, m_cube):
    rows, cols, bands = hyper_cube.shape
    N = rows * cols  # Total number of pixels

    # Initialize the covariance matrix
    covariance_matrix = np.zeros((bands, bands))

    # Reshape the hypercube for easier computation
    reshaped_cube = hyper_cube.reshape(-1, bands)
    # Reshape the mean cube for easier computation
    m_reshaped = m_cube.reshape(-1, bands)
    # Subtract the mean from each pixel spectrum
    mean_subtracted = reshaped_cube - m_reshaped

    # Compute the covariance matrix (outer product of mean-subtracted spectra)
    # Utilize matrix multiplication for efficient computation
    covariance_matrix = np.dot(mean_subtracted.T, mean_subtracted) / N

    return covariance_matrix


def load_cube(hdr_path, dat_path):
    header_file = str(hdr_path)
    spectral_file = str(dat_path)
    numpy_ndarr = envi.open(header_file, spectral_file)
    cube = numpy_ndarr.load().astype('float64')  # Load the entire cube as a numpy array
    return cube


def compute_m(cube):
    num_lines, num_samples, num_bands = cube.shape

    # Define a 3x3 kernel with all ones and then zero the center
    # This kernel will sum up the values of neighbors excluding the center
    kernel = np.ones((3, 3))
    kernel[1, 1] = 0

    # Initialize the result array
    m_res = np.zeros_like(cube)

    # Apply convolution for each band separately
    for c in range(num_bands):
        # Compute the sum of neighbor pixels
        neighbor_sum = convolve(cube[:, :, c], kernel, mode='constant', cval=0.0)

        # Compute the number of valid neighbors for each pixel
        # To handle edges correctly, where neighbors can be outside the image
        valid_neighbor_count = convolve(np.ones_like(cube[:, :, c]), kernel, mode='constant', cval=0.0)

        # Compute the mean by dividing the sum by the count
        # Avoid division by zero by using np.where to leave zeros where the count is zero
        m_res[:, :, c] = np.where(valid_neighbor_count > 0, neighbor_sum / valid_neighbor_count, cube[:, :, c])

    return m_res


def plot_spectrum_graph(spectrum):
    # Plotting the spectrum graph
    plt.figure(figsize=(10, 6))
    plt.plot(range(len(spectrum)), spectrum)
    plt.xlabel('Channel')
    plt.ylabel('Spectral Response')
    plt.title(f'Spectrum Graph of the Target Pixel - {target_pixel_str}')
    plt.grid(True)
    plt.savefig(f"Target_pixel_spec{target_pixel_str}.png", bbox_inches='tight')
    plt.close()  # Close the plot to avoid displaying it in the Python output
    # Add the plot image to the Word document
    doc.add_picture(f"Target_pixel_spec{target_pixel_str}.png", width=Inches(6))


def add_table(NT, WT):
    # Define the table data as a list of lists
    if len(NT.shape) < 3:
        table_data = [
            ["Pixel location", "NT", "WT"],
            ["(5,5)", NT[4, 4], WT[4, 4]],
            ["(10,10)", NT[9, 9], WT[9, 9]],
            ["(15,15)", NT[14, 14], WT[14, 14]],
            ["(20,20)", NT[19, 19], WT[19, 19]]
        ]
    else:
        table_data = [
            ["Pixel location", "NT", "WT"],
            ["(5,5)", NT[4, 4, 0], WT[4, 4, 0]],
            ["(10,10)", NT[9, 9, 0], WT[9, 9, 0]],
            ["(15,15)", NT[14, 14, 0], WT[14, 14, 0]],
            ["(20,20)", NT[19, 19, 0], WT[19, 19, 0]]
        ]

    table = doc.add_table(rows=1, cols=len(table_data[0]))
    table.style = 'Table Grid'

    # Populate the header row
    hdr_cells = table.rows[0].cells
    for i, heading in enumerate(table_data[0]):
        hdr_cells[i].text = heading

    # Add the rest of the data to the table
    for row_data in table_data[1:]:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)


def histogram_plot(NT_results, WT_results, string_1, string2, bins):
    # Calculate histograms without plotting them
    axis = np.linspace(-1000, 1500, bins+1)
    ACE_axis = np.linspace(-150, 500, bins+1)
    if "ACE" in string_1:
        NT_counts, NT_bins = np.histogram(NT_results, ACE_axis)
        WT_counts, WT_bins = np.histogram(WT_results, ACE_axis)
        # Plot the bin centers against the counts
        plt.figure(figsize=(10, 6))
        plt.plot(ACE_axis[:-1], NT_counts, label='NT')
        plt.plot(ACE_axis[:-1], WT_counts, label='WT')
        plt.xlabel('Value')
        plt.ylabel('Density')
        plt.title('Histogram of {} and {}'.format(string_1, string2))
        plt.legend()
        plt.savefig(f"plot{target_pixel_str}.png", bbox_inches='tight')
        plt.close()  # Close the plot to avoid displaying it in the Python output
        # Add the plot image to the Word document
        doc.add_picture(f"plot{target_pixel_str}.png", width=Inches(6))
    else:
        NT_counts, NT_bins = np.histogram(NT_results, axis)
        WT_counts, WT_bins = np.histogram(WT_results, axis)
        # Plot the bin centers against the counts
        plt.figure(figsize=(10, 6))
        plt.plot(axis[:-1], NT_counts, label='NT')
        plt.plot(axis[:-1], WT_counts, label='WT')
        plt.xlabel('Value')
        plt.ylabel('Density')
        plt.title('Continuous Histogram of {} and {}'.format(string_1, string2))
        plt.legend()
        plt.savefig(f"plot{target_pixel_str}.png", bbox_inches='tight')
        plt.close()  # Close the plot to avoid displaying it in the Python output
        # Add the plot image to the Word document
        doc.add_picture(f"plot{target_pixel_str}.png", width=Inches(6))
    return NT_counts, WT_counts


def ace_algorithm(cube, t_T, invers_cov, x_m):

    rows, cols, bands = cube.shape
    ace_scores = np.zeros((rows, cols))

    # Flatten t_T if it's not a 1D array already
    if t_T.ndim > 1:
        t_T = t_T.flatten()

    # Reshape x_m for matrix multiplication, if necessary
    if x_m.shape != cube.shape:
        x_m = x_m.reshape(cube.shape)

    # Compute the part of the ACE algorithm that doesn't change per pixel
    t_t_invers_cov = np.dot(t_T, invers_cov)

    # Calculate the denominator for the ACE algorithm
    denominator_t = np.sqrt(np.dot(t_t_invers_cov, t_T.T))


    for i in range(rows):
        for j in range(cols):
            # Select the pixel spectrum
            pixel_spectrum = x_m[i, j, :]

            # Calculate the numerator for the ACE algorithm
            numerator = np.dot(t_t_invers_cov, pixel_spectrum)

            # Calculate the denominator for the current pixel spectrum
            denominator_x = np.sqrt(np.dot(np.dot(pixel_spectrum.T, invers_cov), pixel_spectrum))

            # Calculate ACE score for the pixel
            if denominator_t > 0 and denominator_x > 0:
                ace_scores[i, j] = numerator / (denominator_t * denominator_x)
            else:
                ace_scores[i, j] = 0  # Handle the case of zero denominators

    return ace_scores


def match_filter(t_i_NT, t_i_WT,inv_covariance_matrix):
    MatchFilter_NT = np.zeros((rows, col))
    MatchFilter_WT = np.zeros((rows, col))
    temp = np.dot(t_T, inv_covariance_matrix)
    for i in range(rows):
        for j in range(col):
            MatchFilter_NT[i, j] = np.dot(temp, t_i_NT[i, j, :])
            MatchFilter_WT[i, j] = np.dot(temp, t_i_WT[i, j, :])

    return MatchFilter_NT, MatchFilter_WT


# File paths
hdr_path = "converted_cube.hdr"
dat_path = "converted_cube.img"

# Path to the saved mean pixels file
filename = 'mean_pixels.pkl'

# Load the mean pixels from the file
with open(filename, 'rb') as f:
    mean_pixels = pickle.load(f)

# Create a new Word Document
doc = Document()

# Load the data cube
cube = load_cube(hdr_path, dat_path)
rows, col, bands = cube.shape

# targets = [(743,16), (485, 25), (652, 97), (410, 136), (740, 23)]

targets = [(410, 136)]
# Calculate the m_vector if it's not already saved
m_vector_filepath = 'm_vector.npy'

# Check if m_vector file doesn't exist
if not os.path.exists(m_vector_filepath):
    print("Calculating the m_vector...")
    m_vector = compute_m(cube)
    np.save(m_vector_filepath, m_vector)  # Save the m_vector to a file for later use
    print(f"m_vector saved to {m_vector_filepath}")
else:
    print(f"m_vector already exists. Loading from {m_vector_filepath}")
    m_vector = np.load(m_vector_filepath)  # Load the m_vector from the file

# p = 0.01  # Target strength
Target_strength = [0.015, 0.02]
# calculate the target implant matrix
t_i_NT_filepath = 't_i_NT.npy'  # Replace with your desired file path

# Check if target implant matrix file doesn't exist
if not os.path.exists(t_i_NT_filepath):  # Check if the file doesn't exist
    print("Calculating the target implant matrix...")
    t_i_NT = cube - m_vector  # x' = x - m
    np.save(t_i_NT_filepath, t_i_NT)  # Save the target implant to a file for later use
    print(f"target implant matrix saved to {t_i_NT_filepath}")
else:
    print(f"target implant matrix already exists. Loading from {t_i_NT_filepath}")
    t_i_NT = np.load(t_i_NT_filepath)  # Load the m_vector from the file
print("calculate the target implant matrix")
for p in Target_strength:
    doc.add_paragraph(f"Target strength {str(p)}")
    for i in range(len(mean_pixels)):

        # Target pixel location (2, 4) in zero-indexed coordinates is (4, 2) in (lines, samples)
        # target_pixel_location = targets[i]
        # target_pixel_str = str(target_pixel_location)
        # doc.add_paragraph(f"Target pixel {target_pixel_str}")
        # target_pixel_spectrum = cube[target_pixel_location[0], target_pixel_location[1], :]
        target_pixel_spectrum = mean_pixels[i]
        target_pixel_str = f"mean pixel in segment {str(i+1)}"
        # Plot the spectrum of the target pixel
        plot_spectrum_graph(target_pixel_spectrum)


        # # Corrupt target pixel
        # target_pixel_str = f"corrupt_spectrum_target_{str(targets[i])}"
        # plot_spectrum_graph(target_pixel_spectrum)
        # target_pixel_spectrum = corrupt_spectrum(target_pixel_spectrum)
        # plot_spectrum_graph(target_pixel_spectrum)
        #
        # # Smoothing corrupted pixel
        # target_pixel_spectrum = fix_corrupted_spectrum(target_pixel_spectrum)
        # target_pixel_str = "Corrupted target after smoothing technique"
        # plot_spectrum_graph(target_pixel_spectrum)


        # x' = x - m +p*t
        t_i_WT = t_i_NT + p * target_pixel_spectrum

        # Reshape target_implant_NT to a 2D array where each row is a pixel and each column is a band
        # target_implant_2D = t_i_NT.reshape(-1, t_i_NT.shape[2])

        # Compute the covariance matrix across all pixels (each pixel is an observation)
        covariance_matrix = compute_cov_matrix(cube, m_vector)
        inv_covariance_matrix = np.linalg.inv(covariance_matrix)
        t_T = target_pixel_spectrum.transpose()
        temp = np.dot(t_T, inv_covariance_matrix)


        MatchFilter_NT = np.zeros((rows,col))
        MatchFilter_WT = np.zeros((rows,col))
        ACE_NT = np.zeros((rows,col))
        ACE_WT = np.zeros((rows,col))

        for i in range(rows):
            for j in range(col):
                MatchFilter_NT[i,j] = np.dot(temp, t_i_NT[i, j, :])
                MatchFilter_WT[i,j] = np.dot(temp, t_i_WT[i, j, :])
        ACE_WT = ace_algorithm(cube, t_T, inv_covariance_matrix, t_i_WT)
        ACE_NT = ace_algorithm(cube, t_T, inv_covariance_matrix, t_i_NT)



        # # save the tables in choosen coordinates
        # doc.add_paragraph("Values of (x'-m) for the following pixels locations, for the first band:")
        # add_table(t_i_NT, t_i_WT)
        # doc.add_paragraph("Match Filter:")
        # add_table(MatchFilter_NT, MatchFilter_WT)
        # doc.add_paragraph("ACE Filter:")
        # add_table(ACE_NT, ACE_WT)

        # Flatten the MatchFilter_NT and MatchFilter_WT arrays to 1D arrays for histogram plotting
        MatchFilter_NT_flat = MatchFilter_NT.flatten()
        MatchFilter_WT_flat = MatchFilter_WT.flatten()
        ACE_WT_flat = ACE_WT.flatten()
        ACE_NT_flat = ACE_NT.flatten()
        bins = 100
        # Histogram of Match Filter with/without target.
        MF_NT_counts, MF_WT_counts = histogram_plot(MatchFilter_NT_flat, MatchFilter_WT_flat, "Match Filter no target", "Match Filter with target", bins)
        ACE_NT_counts, ACE_WT_counts = histogram_plot(ACE_NT_flat, ACE_WT_flat, "ACE Filter no target", "ACE Filter with target", bins)

        X, Y, n_bands = cube.shape
        x_size = len(MatchFilter_WT)
        y_size = len(MatchFilter_NT)

        axs = np.linspace(-1000, 1500, bins+1)

        Pd_MF = 1 - integrate.cumulative_trapezoid(MF_WT_counts, axs[:-1], initial=0) / (25 * X * Y)
        Pfa_MF = 1 - integrate.cumulative_trapezoid(MF_NT_counts, axs[:-1], initial=0) / (25 * X * Y)
        Pd_ACE = 1 - integrate.cumulative_trapezoid(ACE_WT_counts, axs[:-1], initial=0) / (25 * X * Y)
        Pfa_ACE = 1 - integrate.cumulative_trapezoid(ACE_NT_counts, axs[:-1], initial=0) / (25 * X * Y)

        # Plot ACE Inverse Cumulative Probability Distribution
        fig, ax = plt.subplots()
        plt.plot(axs[:-1], Pd_ACE, label='ACE Pd')
        plt.plot(axs[:-1], Pfa_ACE, label='ACE Pfa')
        plt.title('Inverse Cumulative Probability Distribution - ACE')
        plt.legend()
        plt.grid(True)
        plt.savefig("Inverse Cumulative_Probability_Distribution_ACE.png", bbox_inches='tight')
        plt.close()  # Close the plot to avoid displaying it in the Python output
        # Add the plot image to the Word document
        doc.add_picture("Inverse Cumulative_Probability_Distribution_ACE.png", width=Inches(6))

        # Plot MF Inverse Cumulative Probability Distribution
        plt.figure(figsize=(10, 6))
        plt.plot(axs[:-1], Pd_MF, label='MF Pd')
        plt.plot(axs[:-1], Pfa_MF, label='MF Pfa')
        plt.title('Inverse Cumulative Probability Distribution - MF')
        plt.legend()
        plt.grid(True)
        plt.savefig("Inverse_Cumulative_Probability_Distribution_MF.png", bbox_inches='tight')
        plt.close()  # Close the plot to avoid displaying it in the Python output
        # Add the plot image to the Word document
        doc.add_picture("Inverse_Cumulative_Probability_Distribution_MF.png", width=Inches(6))

        # Plot the ROC curves
        plt.figure(figsize=(10, 6))
        plt.plot(Pfa_MF, Pd_MF, label='ROC curve MF')
        plt.plot(Pfa_ACE, Pd_ACE, label='ROC curve ACE')
        plt.plot([0, 1], [0, 1], 'k--')  # Dashed diagonal
        plt.xlim([0, 0.1])
        plt.ylim([0, 1])
        plt.xlabel('Pfa')
        plt.ylabel('Pd')
        plt.title(f"ROC Curve")
        plt.legend()
        plt.grid(True)
        plt.savefig(f"plot{target_pixel_str}.png", bbox_inches='tight')
        plt.close()  # Close the plot to avoid displaying it in the Python output
        # Add the plot image to the Word document
        doc.add_picture(f"plot{target_pixel_str}.png", width=Inches(6))

        # Save the Word document
doc.save('run_3_4.docx')

